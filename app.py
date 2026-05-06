"""
app.py — Backend FastAPI para convertir Markdown a .docx
Uso: uvicorn app:app --reload --host 0.0.0.0 --port 5000
Requiere el entorno Conda de environment.yml
"""

import io
import re
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Inicialización de la app
# ---------------------------------------------------------------------------
app = FastAPI(
    title="Markdown → DOCX Converter",
    description="Convierte texto Markdown a archivos .docx con estilos nativos de Word.",
    version="1.0.0",
)

# CORS: permite peticiones desde el archivo HTML local (file://) y cualquier origen
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],      # En producción restringe a tu dominio
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# Modelo de entrada (Pydantic valida automáticamente)
# ---------------------------------------------------------------------------
class MarkdownPayload(BaseModel):
    markdown: str

# ---------------------------------------------------------------------------
# Helper: detección de secuencias emoji completas
# ---------------------------------------------------------------------------
#
# El orden de las alternativas importa: las secuencias más largas primero.
#
# 1. Keycap sequences      — [0-9#*] + U+FE0F + U+20E3
#    Ejemplos: 4️⃣  #️⃣  *️⃣
#
# 2. Flag sequences (RIS)  — dos Regional Indicator Symbols seguidos
#    Ejemplos: 🇪🇸  🇺🇸  🇦🇷
#
# 3. ZWJ sequences         — emoji + U+200D + emoji (familia, profesión, etc.)
#    Ejemplos: 👨‍💻  👩‍❤️‍👨
#    Se modelan como: (base)(VS16?)(ZWJ + base + VS16?)* (skin tone?)
#
# 4. Emoji + variation selector U+FE0F + U+20E3 (keycap sin dígito base ASCII)
#
# 5. Emoji base + modificador de tono de piel (U+1F3FB–U+1F3FF)
#
# 6. Emoji base + variation selector U+FE0F
#
# 7. Cualquier carácter emoji/símbolo suelto (rangos BMP y suplementarios)

_EMOJI_SEQUENCE_RE = re.compile(
    # — 1. Keycap: dígito/# /* ASCII + FE0F (opcional) + 20E3 —
    r"[0-9#*]\uFE0F?\u20E3"
    # — 2. Banderas: dos Regional Indicator Symbols —
    r"|[\U0001F1E0-\U0001F1FF]{2}"
    # — 3. Secuencias ZWJ (hasta 6 componentes para cubrir familia extendida) —
    r"|(?:[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U0001F004\U0001F0CF]"
    r"\uFE0F?(?:\U0001F3FB|\U0001F3FC|\U0001F3FD|\U0001F3FE|\U0001F3FF)?"
    r"(?:\u200D[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U0001F004\U0001F0CF]"
    r"\uFE0F?(?:\U0001F3FB|\U0001F3FC|\U0001F3FD|\U0001F3FE|\U0001F3FF)?){1,5})"
    # — 4. Emoji base + tono de piel —
    r"|[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U0001F004\U0001F0CF]"
    r"(?:\U0001F3FB|\U0001F3FC|\U0001F3FD|\U0001F3FE|\U0001F3FF)"
    # — 5. Emoji base + variation selector FE0F —
    r"|[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U0001F004\U0001F0CF]\uFE0F"
    # — 6. Símbolos BMP con variation selector —
    r"|[\u2300-\u27BF\u2702-\u27B0\u24C2-\u2BFF]\uFE0F"
    # — 7. Cualquier carácter emoji suelto (suplementarios + BMP) —
    r"|[\U0001F300-\U0001FAFF]"
    r"|[\U0001F004\U0001F0CF\U0001F18F\U0001F191-\U0001F19A"
    r"\U0001F1E6-\U0001F1FF\U0001F201\U0001F202\U0001F21A\U0001F22F"
    r"\U0001F232-\U0001F23A\U0001F250\U0001F251]"
    r"|[\u2600-\u27BF\u2B00-\u2BFF\u3030\u303D\u3297\u3299"
    r"\u231A\u231B\u23E9-\u23F3\u23F8-\u23FA"
    r"\u25AA\u25AB\u25B6\u25C0\u25FB-\u25FE"
    r"\u2614\u2615\u2648-\u2653\u267F\u2693\u26A1\u26AA\u26AB"
    r"\u26BD\u26BE\u26C4\u26C5\u26CE\u26D4\u26EA\u26F2\u26F3"
    r"\u26F5\u26FA\u26FD\u2702\u2705\u2708-\u270D\u270F]",
    flags=re.UNICODE,
)

# Fuentes emoji por orden de preferencia (Word elige la primera disponible)
_EMOJI_FONTS = ['Segoe UI Emoji', 'Apple Color Emoji', 'Noto Color Emoji']


def _set_run_emoji_font(run):
    """
    Asigna fuentes emoji al run usando el XML de rPr directamente,
    porque python-docx no expone w:rFonts con los cuatro atributos a la vez.
    La estructura necesaria en Word es:
        <w:rFonts w:ascii="Segoe UI Emoji"
                  w:hAnsi="Segoe UI Emoji"
                  w:cs="Segoe UI Emoji"
                  w:eastAsia="Segoe UI Emoji"/>
    """
    primary = _EMOJI_FONTS[0]
    rPr = run._r.get_or_add_rPr()

    # Eliminar cualquier w:rFonts previo para no duplicar
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)

    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), primary)
    # Insertar al inicio de rPr (debe ir antes de w:b, w:i, etc.)
    rPr.insert(0, rFonts)


def _add_text_runs(paragraph, text: str, bold=False, italic=False,
                   font_name=None, font_size=None, color=None):
    """
    Añade texto a un párrafo dividiendo la cadena en segmentos normales
    y secuencias emoji completas (keycap, ZWJ, banderas, etc.).
    Los segmentos emoji reciben fuente Segoe UI Emoji;
    los normales heredan los parámetros de formato recibidos.
    """
    # re.split con grupo capturador conserva los separadores en la lista
    parts = _EMOJI_SEQUENCE_RE.split(text)
    emoji_matches = _EMOJI_SEQUENCE_RE.findall(text)

    # Reconstruir el orden intercalado: [normal, emoji, normal, emoji, ...]
    interleaved: list[tuple[str, bool]] = []
    for idx, seg in enumerate(parts):
        if seg:
            interleaved.append((seg, False))
        if idx < len(emoji_matches):
            interleaved.append((emoji_matches[idx], True))

    for chunk, is_emoji in interleaved:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold   = bold
        run.italic = italic
        if is_emoji:
            _set_run_emoji_font(run)
        else:
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if color:
                run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Helpers de formato inline (negrita, cursiva, código inline)
# ---------------------------------------------------------------------------

def apply_inline_formats(paragraph, text: str):
    """
    Parsea el texto buscando **negrita**, *cursiva*, `código` y combinaciones.
    Dentro de cada segmento llama a _add_text_runs para manejar emojis
    correctamente (fuente Segoe UI Emoji) sin romper el resto del formato.
    """
    pattern = re.compile(
        r'(\*\*\*(?P<bold_italic>.+?)\*\*\*)'   # ***bold+italic***
        r'|(\*\*(?P<bold>.+?)\*\*)'              # **bold**
        r'|(\*(?P<italic>.+?)\*)'                # *italic*
        r'|(`(?P<code>.+?)`)',                   # `code`
        re.DOTALL
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Texto plano antes del match (puede contener emojis)
        if m.start() > last_end:
            _add_text_runs(paragraph, text[last_end:m.start()])

        if m.group('bold_italic'):
            _add_text_runs(paragraph, m.group('bold_italic'), bold=True, italic=True)
        elif m.group('bold'):
            _add_text_runs(paragraph, m.group('bold'), bold=True)
        elif m.group('italic'):
            _add_text_runs(paragraph, m.group('italic'), italic=True)
        elif m.group('code'):
            _add_text_runs(
                paragraph, m.group('code'),
                font_name='Courier New',
                font_size=Pt(10),
                color=RGBColor(0xC0, 0x39, 0x2B),
            )

        last_end = m.end()

    # Resto del texto tras el último match
    if last_end < len(text):
        _add_text_runs(paragraph, text[last_end:])


# ---------------------------------------------------------------------------
# Helper: sombreado de celda (XML directo, python-docx no lo expone aún)
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, fill_hex: str):
    """Aplica color de fondo a una celda de tabla vía XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def _set_cell_border(cell, **edges):
    """
    Añade bordes a una celda.
    edges: top, bottom, left, right → dict con 'val', 'sz', 'color'
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, style in edges.items():
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   style.get('val',   'single'))
        el.set(qn('w:sz'),    style.get('sz',    '4'))
        el.set(qn('w:color'), style.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# Helper: bloque de código con estilo visual mejorado
# ---------------------------------------------------------------------------

# Colores de tokens para lenguajes comunes (resaltado básico por regex)
_TOKEN_PATTERNS = [
    # strings dobles y simples
    ('string',   re.compile(r'("(?:[^"\\]|\\.)*"|\'(?:[^\'\\]|\\.)*\')')),
    # comentarios de línea
    ('comment',  re.compile(r'(#.*?$|//.*?$)', re.MULTILINE)),
    # palabras clave Python / JS / common
    ('keyword',  re.compile(
        r'\b(def|class|return|import|from|as|if|elif|else|for|while|in|not|and|or|'
        r'True|False|None|pass|break|continue|raise|try|except|finally|with|yield|'
        r'async|await|lambda|global|nonlocal|del|assert|'
        r'function|var|let|const|new|this|typeof|instanceof|'
        r'public|private|protected|static|void|int|str|float|bool|list|dict|set|'
        r'SELECT|FROM|WHERE|INSERT|UPDATE|DELETE|CREATE|DROP|TABLE|INDEX|JOIN|ON|AS)\b'
    )),
    # números
    ('number',   re.compile(r'\b(\d+\.?\d*)\b')),
]

_TOKEN_COLORS = {
    'string':  RGBColor(0x6A, 0x99, 0x55),   # verde
    'comment': RGBColor(0x6A, 0x6A, 0x6A),   # gris
    'keyword': RGBColor(0x56, 0x9C, 0xD6),   # azul
    'number':  RGBColor(0xB5, 0xCE, 0xA8),   # verde claro
}


def _add_code_block(doc, code_text: str, language: str = ''):
    """
    Agrega un bloque de código al documento con:
    - Fondo gris oscuro (#F3F3F3 en modo claro → se ve bien impreso)
    - Fuente Courier New 9 pt
    - Resaltado de sintaxis básico por regex (strings, keywords, números, comentarios)
    - Etiqueta de lenguaje en la cabecera si se especifica
    """
    # — Cabecera con nombre del lenguaje —
    if language:
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(6)
        hdr.paragraph_format.space_after  = Pt(0)
        hdr.paragraph_format.left_indent  = Inches(0.15)
        run = hdr.add_run(f' {language.upper()} ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(7.5)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # fondo azul en la celda no es posible inline → usamos texto con color
        run.font.color.rgb = RGBColor(0x56, 0x9C, 0xD6)

    # Cada línea del bloque es un párrafo independiente para resaltado correcto
    first_line = True
    for raw_line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent   = Inches(0.25)
        p.paragraph_format.right_indent  = Inches(0.15)
        p.paragraph_format.space_before  = Pt(0) if not first_line else Pt(4)
        p.paragraph_format.space_after   = Pt(0)
        first_line = False

        # Colorear tokens en la línea
        _apply_syntax_highlight(p, raw_line if raw_line else ' ')

    # Espacio posterior al bloque
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(6)


def _apply_syntax_highlight(paragraph, line: str):
    """
    Divide la línea en tokens y les aplica color.
    Los fragmentos sin match quedan en color base (texto claro sobre fondo oscuro).
    """
    BASE_COLOR = RGBColor(0x1E, 0x1E, 0x1E)   # casi negro → sobre fondo gris claro imprime bien

    # Construir lista de (start, end, token_type) sin solapar
    spans: list[tuple[int, int, str]] = []
    occupied = set()

    for token_type, pattern in _TOKEN_PATTERNS:
        for m in pattern.finditer(line):
            if any(m.start() <= pos < m.end() for pos in occupied):
                continue
            spans.append((m.start(), m.end(), token_type))
            occupied.update(range(m.start(), m.end()))

    spans.sort(key=lambda x: x[0])

    cursor = 0
    for start, end, token_type in spans:
        # texto sin token antes del match
        if cursor < start:
            _add_text_runs(paragraph, line[cursor:start],
                           font_name='Courier New', font_size=Pt(9), color=BASE_COLOR)
        # texto del token
        run = paragraph.add_run(line[start:end])
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = _TOKEN_COLORS[token_type]
        cursor = end

    # resto de la línea
    if cursor < len(line):
        _add_text_runs(paragraph, line[cursor:],
                       font_name='Courier New', font_size=Pt(9), color=BASE_COLOR)


# ---------------------------------------------------------------------------
# Helper: tabla Markdown → tabla Word
# ---------------------------------------------------------------------------

def _add_markdown_table(doc, header_cells: list[str], rows: list[list[str]], alignments: list[str]):
    """
    Construye una tabla Word con:
    - Cabecera fondo azul oscuro, texto blanco, negrita
    - Filas alternas con fondo blanco / gris muy claro
    - Bordes finos en todas las celdas
    - Alineación por columna según la sintaxis Markdown (:---, :---:, ---:)
    """
    n_cols = len(header_cells)
    n_rows = len(rows)

    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    HEADER_BG   = '2E4057'   # azul pizarra
    ROW_EVEN_BG = 'FFFFFF'
    ROW_ODD_BG  = 'F2F4F7'
    BORDER_COLOR = 'C5CCD6'

    border_style = {'val': 'single', 'sz': '4', 'color': BORDER_COLOR}

    # ── Fila de cabecera ──────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for col_idx, cell_text in enumerate(header_cells):
        cell = hdr_row.cells[col_idx]
        _set_cell_shading(cell, HEADER_BG)
        _set_cell_border(cell, top=border_style, bottom=border_style,
                               left=border_style, right=border_style)
        p = cell.paragraphs[0]
        p.alignment = _md_alignment(alignments[col_idx] if col_idx < len(alignments) else '')
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # ── Filas de datos ────────────────────────────────────────────────────
    for row_idx, row_data in enumerate(rows):
        bg = ROW_ODD_BG if row_idx % 2 else ROW_EVEN_BG
        tr = table.rows[row_idx + 1]
        for col_idx in range(n_cols):
            cell_text = row_data[col_idx].strip() if col_idx < len(row_data) else ''
            cell = tr.cells[col_idx]
            _set_cell_shading(cell, bg)
            _set_cell_border(cell, top=border_style, bottom=border_style,
                                   left=border_style, right=border_style)
            p = cell.paragraphs[0]
            p.alignment = _md_alignment(alignments[col_idx] if col_idx < len(alignments) else '')
            apply_inline_formats(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(10)

    # Espacio después de la tabla
    doc.add_paragraph()


def _md_alignment(separator_cell: str):
    """Deduce la alineación Word a partir de la celda separadora de tabla Markdown."""
    s = separator_cell.strip()
    if s.startswith(':') and s.endswith(':'):
        return WD_ALIGN_PARAGRAPH.CENTER
    if s.endswith(':'):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _parse_table_row(line: str) -> list[str]:
    """Divide una línea de tabla Markdown en celdas, eliminando los `|` externos."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return line.split('|')


def _is_separator_row(line: str) -> bool:
    """Detecta si la línea es la fila separadora de cabecera: | --- | :---: | ---: |"""
    return bool(re.match(r'^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?\s*$', line))


# ---------------------------------------------------------------------------
# Conversor principal Markdown → Document
# ---------------------------------------------------------------------------

def markdown_to_docx(md_text: str) -> io.BytesIO:
    """
    Convierte una cadena Markdown a un objeto Document de python-docx
    y lo devuelve como un stream BytesIO listo para descarga.

    Mapeo de estilos:
        # → 'Heading 1'   (Word genera el índice automáticamente)
        ## → 'Heading 2'
        ### → 'Heading 3'
        #### → 'Heading 4'
        - / * / + → lista con viñetas  ('List Bullet')
        1. 2. … → lista numerada       ('List Number')
        > → cita en bloque             (cursiva + sangría)
        ``` lang … ``` → bloque de código con resaltado de sintaxis
        | col | col | → tabla con cabecera estilizada y filas alternas
        texto normal → 'Normal'
    """
    doc = Document()

    # Márgenes del documento (2.5 cm cada lado)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Bloque de código cercado (``` lang) ──────────────────────────
        if line.strip().startswith('```'):
            # Extraer nombre de lenguaje opcional: ```python, ```js, etc.
            lang = line.strip()[3:].strip().lower()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, '\n'.join(code_lines), language=lang)
            i += 1
            continue

        # ── Tabla Markdown ────────────────────────────────────────────────
        # Una tabla empieza cuando la línea actual tiene '|' y la siguiente
        # es la fila separadora (| --- | :---: | ---: |)
        if '|' in line and i + 1 < len(lines) and _is_separator_row(lines[i + 1]):
            header_cells = _parse_table_row(line)
            alignments   = _parse_table_row(lines[i + 1])
            data_rows    = []
            i += 2   # saltar cabecera + separador
            while i < len(lines) and '|' in lines[i]:
                data_rows.append(_parse_table_row(lines[i]))
                i += 1
            _add_markdown_table(doc, header_cells, data_rows, alignments)
            continue

        # ── Encabezados (H1–H6) ──────────────────────────────────────────
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            content = heading_match.group(2).strip()
            safe_level = min(level, 4)
            p = doc.add_heading(level=safe_level)
            apply_inline_formats(p, content)
            i += 1
            continue

        # ── Lista con viñetas ─────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bullet_match:
            content = bullet_match.group(2).strip()
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formats(p, content)
            i += 1
            continue

        # ── Lista numerada ────────────────────────────────────────────────
        numbered_match = re.match(r'^\d+\.\s+(.*)', line)
        if numbered_match:
            content = numbered_match.group(1).strip()
            p = doc.add_paragraph(style='List Number')
            apply_inline_formats(p, content)
            i += 1
            continue

        # ── Cita en bloque ────────────────────────────────────────────────
        blockquote_match = re.match(r'^>\s?(.*)', line)
        if blockquote_match:
            content = blockquote_match.group(1).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
            i += 1
            continue

        # ── Línea horizontal ─────────────────────────────────────────────
        if re.match(r'^[-*_]{3,}\s*$', line):
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # ── Línea vacía ───────────────────────────────────────────────────
        if line.strip() == '':
            i += 1
            continue

        # ── Párrafo normal ────────────────────────────────────────────────
        p = doc.add_paragraph(style='Normal')
        apply_inline_formats(p, line.strip())
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Rutas FastAPI
# ---------------------------------------------------------------------------

@app.get("/", tags=["health"])
async def health_check():
    """Health-check simple."""
    return {"status": "ok", "message": "Markdown → DOCX converter running"}


@app.post("/convert", tags=["converter"])
async def convert(payload: MarkdownPayload):
    """
    Recibe JSON: { "markdown": "..." }
    Devuelve el archivo .docx como stream descargable.
    """
    if not payload.markdown.strip():
        raise HTTPException(
            status_code=422,
            detail="El texto Markdown no puede estar vacío."
        )

    try:
        docx_buffer = markdown_to_docx(payload.markdown)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error interno al generar el documento: {str(e)}"
        )

    headers = {
        "Content-Disposition": 'attachment; filename="documento.docx"'
    }

    # StreamingResponse es el equivalente FastAPI de send_file()
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


# ---------------------------------------------------------------------------
# Punto de entrada (ejecución directa: python app.py)
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=5000, reload=True)

